import json
from datetime import datetime, timedelta

import stripe
from django.conf import settings
from django.contrib import messages
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.mixins import LoginRequiredMixin
from django.core.exceptions import PermissionDenied
from django.core.paginator import Paginator
from django.db.models import Q, Avg
from django.db.models.aggregates import Sum
from django.http import JsonResponse, HttpResponseForbidden, HttpResponse
from django.shortcuts import get_object_or_404, render, redirect
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
from django.views.generic import TemplateView, DetailView, ListView
from docx import Document
from docx.shared import Pt

from .filters import CarFilter
from .forms import CarReviewForm
from .models import Car, Rental, Transaction, TripTracking, CarLocation, Dealer, Favorite, CarReview


def superuser_required(view_func):
	return user_passes_test(lambda u: u.is_superuser)(view_func)


def home(request):
    # Используем ваш способ получения популярных и рекомендованных машин
    # Если '-reviews__rating' вызывает ошибку (например, нет такой связи/поля),
    # замените на другой способ сортировки, например, '-created_at'
    try:
        popular_qs = Car.objects.order_by('-reviews__rating').filter(is_available=True)[:4]
    except Exception: # Откат к простой сортировке, если '-reviews__rating' не работает
        popular_qs = Car.objects.filter(is_available=True).order_by('-created_at')[:4]

    recommended_qs = Car.objects.filter(is_available=True).order_by('?')[:8] # '?' для случайного порядка, если БД поддерживает

    favorite_car_ids = set()
    if request.user.is_authenticated:
        favorite_car_ids = set(
            Favorite.objects.filter(user=request.user).values_list("car_id", flat=True)
        )

    popular_cars_list = list(popular_qs) # Преобразуем queryset в список для добавления атрибута
    for car_obj in popular_cars_list:
        car_obj.is_favorited = car_obj.pk in favorite_car_ids

    recommended_cars_list = list(recommended_qs) # Преобразуем queryset в список
    for car_obj in recommended_cars_list:
        car_obj.is_favorited = car_obj.pk in favorite_car_ids

    context = {
       "popular_cars": popular_cars_list,
       "recommended_cars": recommended_cars_list,
    }
    return render(request, "home.html", context)


def car_list(request):
	# 1) Фильтрация через django-filter
	f = CarFilter(request.GET, queryset=Car.objects.filter(is_available=True))
	qs = f.qs

	# 2) Собираем параметры для формы (могут понадобиться в шаблоне)
	selected_types = request.GET.getlist("type")
	selected_seats = request.GET.get("seats", "")
	selected_dealers = request.GET.getlist("dealer")
	price_to = request.GET.get("price_to", "")
	query = request.GET.get("q", "")

	# 3) Пагинация
	paginator = Paginator(qs, 12)
	page_num = request.GET.get("page")
	cars_page = paginator.get_page(page_num)

	# 4) Избранные авто текущего пользователя
	if request.user.is_authenticated:
		favorite_car_ids = set(
			Favorite.objects
			.filter(user=request.user, car__in=cars_page.object_list)
			.values_list("car_id", flat=True)
		)
	else:
		favorite_car_ids = set()

	for car in cars_page:
		car.is_favorited = (car.pk in favorite_car_ids)

	# 5) Контекст для шаблона
	context = {
		"filter": f,
		"cars": cars_page,
		"selected_types": selected_types,
		"selected_seats": selected_seats,
		"selected_dealers": selected_dealers,
		"price_to": price_to,
		"query": query,
		"params": request.GET.urlencode(),
		"seats_list": [2, 4, 6, 8],
		"favorite_car_ids": favorite_car_ids,
	}

	# 6) Если это HTMX-запрос — возвращаем только карточки
	if request.headers.get("Hx-Request") == "true":
		return render(request, "partials/car_cards.html", context)

	# 7) Иначе — полный рендер с сайдбаром + карточками
	return render(request, "car_list.html", context)


def car_search(request):
	"""
	Страница «/search/» — результаты поиска + тот же сайд‑бар фильтров.
	Достаточно передать тот же CarFilter; параметр «q» уже там обрабатывается.
	"""
	f = CarFilter(request.GET, queryset=Car.objects.filter(is_available=True))
	paginator = Paginator(f.qs, 12)
	cars_page = paginator.get_page(request.GET.get("page"))

	return render(request, "car_search.html", {
		"filter": f,
		"cars": cars_page,
		"query": request.GET.get("q", ""),
	})


@login_required
def favorites(request):
    """
    Страница «Избранное»: показывает список машин, добавленных пользователем в избранное.
    """
    fav_qs = Favorite.objects.filter(user=request.user).select_related('car')
    cars_list = []
    for fav in fav_qs:
        car_obj = fav.car
        car_obj.is_favorited = True # Все машины на этой странице по определению в избранном
        cars_list.append(car_obj)

    # Убедитесь, что имя шаблона "car_favourites.html" (или какое у вас на самом деле) верное
    return render(request, "car_favourites.html", {
       "cars": cars_list,
    })



def faq(request):
	"""
	Статическая страница FAQ по каршерингу
	"""
	return render(request, 'faq.html')


@login_required
@require_POST  # Убеждаемся, что принимаем только POST запросы
@csrf_exempt
def favorite_toggle(request, pk):
    car = get_object_or_404(Car, pk=pk)
    user = request.user

    favorite_instance = Favorite.objects.filter(user=user, car=car).first()
    is_now_favorited: bool

    if favorite_instance:
       favorite_instance.delete()
       is_now_favorited = False
    else:
       Favorite.objects.create(user=user, car=car)
       is_now_favorited = True

    context = {
       'car': car, # car объект нужен для {% url 'favorite_toggle' car.pk %} в fav_button.html
       'favorited': is_now_favorited, # Передаем АКТУАЛЬНОЕ состояние
    }

    return render(request, 'partials/fav_button.html', context)


def car_detail(request, pk):
	car = get_object_or_404(Car, pk=pk)

	# Рейтинг
	# Убедитесь, что 'reviews' это правильный related_name от вашей модели Review к Car
	rating_data = car.reviews.aggregate(avg=Avg("rating"))
	# Если нет отзывов, aggregate вернет {'avg': None}
	rating = rating_data["avg"] if rating_data["avg"] is not None else 0  # Ставим 0 или другое значение по умолчанию

	# «Похожие» по типу, без текущей
	similar = (
		Car.objects
		.filter(car_type=car.car_type, is_available=True)
		.exclude(pk=car.pk)[:6]
	)

	# «Недавние», без текущей
	recent_cars = (
		Car.objects
		.filter(is_available=True)
		.exclude(pk=car.pk)
		.order_by("-id")[:6]  # или -created_at, если есть такое поле
	)

	# «Рекомендуемые» у того же дилера
	recommend_cars = []
	if hasattr(car, 'dealer') and car.dealer:  # Проверка, есть ли у машины дилер
		recommend_cars = (
			Car.objects
			.filter(dealer=car.dealer, is_available=True)
			.exclude(pk=car.pk)[:6]
		)

	# Все отзывы
	reviews = car.reviews.select_related("user").order_by("-created_at")

	# Собираем галерею - ИСПРАВЛЕННАЯ ЛОГИКА
	images = []
	# 1. Добавляем URL главного изображения, если оно есть и у него есть URL
	if car.main_image and hasattr(car.main_image, 'url'):
		images.append(car.main_image.url)

	# 2. Добавляем URL'ы из ManyToMany поля photos
	# car.photos.all() возвращает QuerySet связанных объектов CarPhoto
	# Сортируем согласно Meta.ordering в модели CarPhoto
	additional_photos = car.photos.all().order_by('order', 'uploaded_at')
	for photo_obj in additional_photos:
		# Проверяем, что у объекта CarPhoto есть файл в поле image и у него есть URL
		if photo_obj.image and hasattr(photo_obj.image, 'url'):
			images.append(photo_obj.image.url)

	# Если список изображений пуст (нет ни главного, ни дополнительных),
	# можно добавить URL изображения-заглушки:
	if not images:
		from django.templatetags.static import static  # или {% static 'path/to/placeholder.png' %} в шаблоне
		images.append(static('images/car_placeholder.png'))  # Укажите ваш путь к заглушке

	# Характеристики
	characteristics = []
	if hasattr(car, 'get_car_type_display'):  # Проверка существования метода
		characteristics.append(("Тип", car.get_car_type_display()))
	if hasattr(car, 'get_transmission_display'):
		characteristics.append(("Трансмиссия", car.get_transmission_display()))
	if hasattr(car, 'get_fuel_type_display'):
		characteristics.append(("Топливо", car.get_fuel_type_display()))
	if hasattr(car, 'seats'):
		characteristics.append(("Мест", car.seats))
	if hasattr(car, 'doors'):
		characteristics.append(("Дверей", car.doors))
	if hasattr(car, 'mileage'):
		characteristics.append(("Пробег", f"{car.mileage} км" if car.mileage is not None else "N/A"))
	if hasattr(car, 'air_conditioner'):
		characteristics.append(("Кондиционер", "Да" if car.air_conditioner else "Нет"))
	if hasattr(car, 'year'):
		characteristics.append(("Год выпуска", car.year))

	if hasattr(car, 'features') and car.features:  # Проверяем, что car.features существует и не None
		for key, val in car.features.items():
			characteristics.append((key, "Да" if val is True else str(val)))  # Приводим val к строке для безопасности

	# Фильтр — убираем текущую машину из исходного qs
	# Убедитесь, что CarFilter правильно определен и импортирован
	fil = CarFilter(
	    request.GET,
	    queryset=Car.objects.filter(is_available=True)
	                         .exclude(pk=car.pk)
	)

	print(images)

	context = {
		"car": car,
		"images": images,  # Передаем исправленный список URL
		"filter": fil,
		"rating": int(rating),  # Передаем рейтинг как целое число для цикла в шаблоне
		"similar": similar,
		"recent_cars": recent_cars,
		"recommend_cars": recommend_cars,
		"reviews": reviews,
		"characteristics": characteristics,
	}
	# Убедитесь, что имя шаблона "car_detail.html" указано верно
	return render(request, "car_detail.html", context)


class ContactUsView(TemplateView):
	template_name = "contact.html"


class AboutUsView(TemplateView):
	template_name = "about.html"


class RentalDetailView(LoginRequiredMixin, DetailView):
	model = Rental
	template_name = "rental_detail.html"
	context_object_name = "rental"

	def get_queryset(self):
		# Ограничиваем выборку аренд только текущим пользователем
		return Rental.objects.filter(user=self.request.user)


class TransactionHistoryView(ListView):
	model = Transaction
	template_name = 'payment_history.html'
	context_object_name = 'payments'

	def get_queryset(self):
		return Transaction.objects.filter(user=self.request.user).order_by('-timestamp')


class MyRentalsView(LoginRequiredMixin, TemplateView):
	template_name = 'active_rentals.html'  # Можно назвать, например, my_rentals.html

	def get_context_data(self, **kwargs):
		context = super().get_context_data(**kwargs)
		# Активные аренды – например, аренды, которые оплачены и имеют статус COMPLETED
		context['active_rentals'] = Rental.objects.filter(user=self.request.user, is_paid=True, status="OWNED")
		# История (неактивные аренды) – аренды, которые не имеют статус COMPLETED
		context['inactive_rentals'] = Rental.objects.filter(user=self.request.user).exclude(status="COMPLETED")
		return context


def user_is_retailer(user):
	"""
	Проверяет, является ли пользователь суперпользователем или входит в группу "Retailers"
	"""
	return user.is_superuser or user.groups.filter(name="Retailers").exists()


@login_required
def retailer_tracking_dashboard(request, retailer_id):
	"""
	Дэшборд для просмотра трекинга машин конкретного ритейлера.
	Фильтруется по типу автомобиля и доступности, при этом выбираются
	только те машины, у которых car.dealer_id совпадает с retailer_id.
	"""
	# Получаем фильтры из GET-параметров
	car_type = request.GET.get('car_type')
	is_available = request.GET.get('is_available')

	# Выбираем объекты CarLocation с подгрузкой связанных моделей и фильтруем по дилеру
	locations = CarLocation.objects.select_related('car', 'car__dealer').filter(car__dealer_id=retailer_id)
	if car_type:
		locations = locations.filter(car__type=car_type)
	if is_available:
		if is_available.lower() == 'true':
			locations = locations.filter(car__is_available=True)
		elif is_available.lower() == 'false':
			locations = locations.filter(car__is_available=False)

	# Список возможных типов автомобилей для фильтрации
	car_types = Car.CarType.choices

	context = {
		'locations': locations,
		'car_types': car_types,
		'selected_car_type': car_type,
		'selected_is_available': is_available,
		'retailer_id': retailer_id,
	}
	return render(request, 'retailer_tracking_dashboard.html', context)


@login_required
def retailer_car_tracking_history(request, retailer_id, car_id):
	"""
	Страница истории трекинга для конкретной машины.
	Проверяем, что машина принадлежит ритейлеру с идентификатором retailer_id.
	"""
	car = get_object_or_404(Car, id=car_id, dealer_id=retailer_id)

	# Получаем текущее местоположение машины (если оно имеется)
	try:
		location = car.location
	except CarLocation.DoesNotExist:
		location = None

	# Получаем историю трекинга (все записи по всем арендам данной машины)
	tracking_logs = TripTracking.objects.filter(rental__car=car).order_by('timestamp')

	context = {
		'car': car,
		'location': location,
		'tracking_logs': tracking_logs,
		'retailer_id': retailer_id,
	}
	return render(request, 'retailer_car_tracking_history.html', context)


# === НАЧАЛО НОВОГО КОДА: СКОПИРУЙТЕ ВСЮ ФУНКЦИЮ ===
@login_required
def cancel_rental_view(request, rental_id):
	"""
	Обрабатывает отмену бронирования пользователем.
	"""
	# Ищем аренду, которая принадлежит пользователю и еще не началась
	rental = get_object_or_404(Rental, id=rental_id, user=request.user)

	if request.method == 'POST':
		# 1. Меняем статус аренды на "Отменено"
		rental.status = 'CANCELED'
		rental.save()

		# 2. Освобождаем автомобиль, делаем его снова доступным
		car = rental.car
		car.status = 'available'
		car.save()

		# 3. Сообщаем пользователю об успехе
		messages.success(request, f"Бронирование автомобиля {car.brand} {car.model} было отменено.")

		# 4. Перенаправляем в личный кабинет
		return redirect('dashboard')

	# Если это GET-запрос, показываем страницу подтверждения
	context = {
		'rental': rental
	}
	return render(request, 'cancel_rental_confirmation.html', context)


@login_required
def retailers_dashboard(request):
	# Убедимся, что у пользователя есть профиль дилера
	try:
		dealer = request.user.dealer
	except AttributeError:
		# Если у пользователя нет профиля дилера, можно перенаправить его или показать ошибку
		return redirect('home')

	# Получаем все автомобили, связанные с этим дилером
	all_cars = Car.objects.filter(dealer=dealer)

	# Фильтруем автомобили по статусам
	pending_cars = all_cars.filter(status='pending_inspection')
	available_cars = all_cars.filter(status='available')
	rented_cars = all_cars.filter(status='rented')

	# Получаем активные и завершенные аренды для этого дилера
	active_rentals = Rental.objects.filter(car__dealer=dealer, status='OWNED')
	completed_rentals = Rental.objects.filter(car__dealer=dealer, status='COMPLETED')

	context = {
	    'dealer': dealer,
	    'pending_cars': pending_cars,      # <-- Наш новый список
	    'available_cars': available_cars,
	    'rented_cars': rented_cars,
	    'active_rentals': active_rentals,
	    'completed_rentals': completed_rentals,
	    'total_cars': all_cars.count(),
	}
	return render(request, 'retailers_dashboard.html', context)


@login_required
def confirm_inspection_view(request, car_id):
    # Получаем автомобиль
    car = get_object_or_404(Car, id=car_id)

    # Проверяем, что текущий пользователь является дилером этого автомобиля
    if not hasattr(request.user, 'dealer') or car.dealer != request.user.dealer:
        raise PermissionDenied("У вас нет прав для выполнения этого действия.")

    # Меняем статус только если он 'pending_inspection'
    if car.status == 'pending_inspection':
        car.status = 'available'
        car.save()
        messages.success(request, f'Автомобиль "{car}" снова доступен для аренды.')
    else:
        messages.warning(request, f'Статус автомобиля "{car}" уже был изменен.')

    # Возвращаем дилера в его панель
    return redirect('retailers_dashboard')



@login_required
def finish_rental_view(request, rental_id):
    rental = get_object_or_404(Rental, id=rental_id, user=request.user)
    car = rental.car

    if request.method == 'POST':
        form = CarReviewForm(request.POST)
        if form.is_valid():
            # 1. Сохраняем отзыв
            review = form.save(commit=False)
            review.user = request.user
            review.car = car
            review.save()

            # 2. Обновляем статус аренды
            rental.status = 'COMPLETED'
            rental.save()

            # 3. Обновляем статус автомобиля
            # (Более продвинутый вариант - см. Фичу 2)
            car.status = 'pending_inspection'
            car.save()

            # 4. Перенаправляем пользователя в личный кабинет
            return redirect('dashboard')
    else:
        form = CarReviewForm()

    context = {
        'rental': rental,
        'form': form
    }
    return render(request, 'finish_rental.html', context)


def api_car_tracking(request, retailer_id, car_id):
	"""
	API-эндпоинт для получения данных трекинга конкретной машины,
	принадлежащей ритейлеру с идентификатором retailer_id.
	"""
	car = get_object_or_404(Car, id=car_id, dealer_id=retailer_id)
	try:
		location = car.location
		loc_data = {
			'latitude': float(location.latitude),
			'longitude': float(location.longitude),
			'updated_at': location.updated_at.isoformat(),
		}
	except CarLocation.DoesNotExist:
		loc_data = None

	tracking_logs = TripTracking.objects.filter(rental__car=car).order_by('timestamp')
	logs_data = [
		{
			'latitude': float(log.latitude),
			'longitude': float(log.longitude),
			'timestamp': log.timestamp.isoformat(),
		}
		for log in tracking_logs
	]
	return JsonResponse({
		'location': loc_data,
		'tracking_logs': logs_data,
	})


def api_add_car_tracking(request, retailer_id, car_id):
	if request.method != "POST":
		return JsonResponse({'error': 'Метод не разрешен. Используйте POST.'}, status=405)

	# Проверяем, что машина принадлежит ритейлеру (dealer)
	car = get_object_or_404(Car, id=car_id, dealer_id=retailer_id)

	# Получаем rental_id из POST-данных
	rental_id = request.POST.get('rental_id')
	if not rental_id:
		return JsonResponse({'error': 'Параметр rental_id обязателен.'}, status=400)

	try:
		rental_id = int(rental_id)
	except ValueError:
		return JsonResponse({'error': 'rental_id должен быть числовым.'}, status=400)

	# Проверяем, что аренда существует для этой машины
	try:
		rental = car.rental_set.get(id=rental_id)
	except Rental.DoesNotExist:
		return JsonResponse({'error': 'Аренда с данным rental_id не найдена для этой машины.'}, status=404)

	# Получаем координаты из POST-данных
	latitude = request.POST.get('latitude')
	longitude = request.POST.get('longitude')

	if latitude is None or longitude is None:
		return JsonResponse({'error': 'Параметры latitude и longitude обязательны.'}, status=400)

	try:
		latitude = float(latitude)
		longitude = float(longitude)
	except ValueError:
		return JsonResponse({'error': 'Неверный формат координат.'}, status=400)

	# Создаем новую запись трекинга
	tracking = TripTracking.objects.create(
		rental=rental,
		latitude=latitude,
		longitude=longitude
	)

	return JsonResponse({
		'status': 'success',
		'tracking': {
			'id': tracking.id,
			'latitude': float(tracking.latitude),
			'longitude': float(tracking.longitude),
			'timestamp': tracking.timestamp.isoformat(),
		}
	})


stripe.api_key = settings.STRIPE_SECRET_KEY


@login_required
def checkout_view(request, car_id):
	"""
	Создаёт запись Rental с placeholder-данными (is_paid=False) и передаёт rental_id в шаблон.
	При дальнейшей обработке заполнение данных можно обновить.
	"""
	car = get_object_or_404(Car, id=car_id)

	# Заполняем placeholder-данными необходимые поля
	now = datetime.now()
	today = now.date()
	dummy_time = now.time().replace(hour=0, minute=0, second=0, microsecond=0)

	rental = Rental.objects.create(
		user=request.user,
		car=car,
		start_time=now,
		end_time=now + timedelta(hours=24),  # placeholder – аренда на 1 день
		full_name=request.user.get_full_name() or request.user.username,
		phone_number="",
		address="",
		city="",
		pickup_location="",
		pickup_date=today,
		pickup_time=dummy_time,
		dropoff_location="",
		dropoff_date=today,
		dropoff_time=dummy_time,
		payment_method="CARD",
		is_paid=False,
		total_price=car.price_per_day  # placeholder, можно пересчитать по нужной логике
	)

	return render(request, "rent_car.html", {
		"car": car,
		"stripe_pk": settings.STRIPE_PUBLISHABLE_KEY,
		"rental_id": rental.id,
	})


@csrf_exempt
@login_required
def update_rental(request):
	"""
	Обновляет данные Rental, полученные из формы.
	Ожидается POST-запрос с данными и rental_id.
	"""
	if request.method == "POST":
		rental_id = request.POST.get("rental_id")
		rental = get_object_or_404(Rental, id=rental_id, user=request.user)
		rental.full_name = request.POST.get("full_name", rental.full_name)
		rental.phone_number = request.POST.get("phone_number", rental.phone_number)
		rental.address = request.POST.get("address", rental.address)
		rental.city = request.POST.get("city", rental.city)
		rental.pickup_location = request.POST.get("pickup_location", rental.pickup_location)
		rental.pickup_date = request.POST.get("pickup_date", rental.pickup_date)
		rental.pickup_time = request.POST.get("pickup_time", rental.pickup_time)
		rental.dropoff_location = request.POST.get("dropoff_location", rental.dropoff_location)
		rental.dropoff_date = request.POST.get("dropoff_date", rental.dropoff_date)
		rental.dropoff_time = request.POST.get("dropoff_time", rental.dropoff_time)
		rental.save()
		return JsonResponse({"status": "success"})
	return JsonResponse({"error": "Метод не разрешен"}, status=405)


@csrf_exempt
def create_payment_intent(request):
	"""
	Создаёт PaymentIntent и передаёт rental_id в metadata.
	Ожидает JSON с полями: amount (сумма в USD, например, цена за день) и rental_id.
	"""
	data = json.loads(request.body)
	amount = data.get("amount")
	rental_id = data.get("rental_id")

	intent = stripe.PaymentIntent.create(
		amount=int(amount * 100),  # переводим сумму в центы
		currency="KZT",
		automatic_payment_methods={"enabled": True},
		metadata={"rental_id": rental_id}
	)
	return JsonResponse({"clientSecret": intent.client_secret})


@login_required
def processing_view(request):
	rental_id = request.GET.get("rental_id")
	if not rental_id:
		return HttpResponseForbidden("Rental ID не передан")
	rental = get_object_or_404(Rental, id=rental_id, user=request.user)
	return render(request, "processing.html", {"rental": rental})


@login_required
def check_transaction(request):
	"""
	Возвращает статус аренды по rental_id: "success" если rental.is_paid True, иначе "processing".
	"""
	rental_id = request.GET.get("rental_id")
	if not rental_id:
		return JsonResponse({"error": "Rental ID required"}, status=400)

	rental = get_object_or_404(Rental, id=rental_id, user=request.user)
	if settings.DEBUG:
		return JsonResponse({"status": "success"})
	if rental.is_paid:
		return JsonResponse({"status": "success"})
	return JsonResponse({"status": "processing"})


# Страница подтверждения аренды с информацией о машине, дилере и картой (центр карты — Астана)
@login_required
def rental_success(request):
	rental_id = request.GET.get("rental_id")
	if not rental_id:
		return HttpResponseForbidden("Rental ID не передан")
	rental = get_object_or_404(Rental, id=rental_id, user=request.user)
	if settings.DEBUG:
		rental.is_paid = True
		rental.status = "OWNED"
		rental.save()

	return render(request, "rental_success.html", {"rental": rental})


@login_required
def download_receipt_docx(request):
	import io

	rental_id = request.GET.get("rental_id")
	if not rental_id:
		return HttpResponseForbidden("Rental ID не передан")
	rental = get_object_or_404(Rental, id=rental_id, user=request.user)

	# 1) Создаем документ
	doc = Document()
	style = doc.styles["Normal"]
	style.font.name = "Arial"
	style.font.size = Pt(12)

	# Заголовок
	h = doc.add_heading("Квитанция по аренде", level=1)
	h.alignment = 1  # по центру

	# Таблица с данными
	table = doc.add_table(rows=0, cols=2)
	table.style = "Table Grid"

	def add_row(label, value):
		row_cells = table.add_row().cells
		row_cells[0].text = str(label)
		row_cells[1].text = str(value)

	add_row("ID аренды", rental.id)
	add_row("Пользователь", rental.user.get_full_name() or rental.user.username)
	add_row("Автомобиль", f"{rental.car.brand} {rental.car.model}")
	dealer = rental.car.dealer.name if rental.car.dealer else "Не указан"
	add_row("Дилер", dealer)
	add_row("Дата получения", f"{rental.pickup_date} {rental.pickup_time}")
	add_row("Дата возврата", f"{rental.dropoff_date} {rental.dropoff_time}")
	add_row("Итоговая цена (₸)", rental.total_price)
	add_row("Оплачено", "Да" if rental.is_paid else "Нет")

	# 2) Сохраняем в память
	buffer = io.BytesIO()
	doc.save(buffer)
	buffer.seek(0)

	# 3) Отдаем как скачиваемый файл
	response = HttpResponse(
		buffer.read(),
		content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
	)
	response["Content-Disposition"] = (
		f'attachment; filename="receipt_rental_{rental.id}.docx"'
	)
	return response


@csrf_exempt
def stripe_webhook(request):
	"""
	Обрабатывает события Stripe.
	При событии payment_intent.succeeded обновляет rental.is_paid на True.
	При payment_intent.payment_failed оставляет is_paid=False.
	"""
	payload = request.body
	sig_header = request.META.get("HTTP_STRIPE_SIGNATURE")
	webhook_secret = settings.STRIPE_WEBHOOK_SECRET  # настройте в settings.py
	try:
		event = stripe.Webhook.construct_event(payload, sig_header, webhook_secret)
	except Exception as e:
		return HttpResponse(status=400)

	if event["type"] == "payment_intent.succeeded":
		intent = event["data"]["object"]
		rental_id = intent.get("metadata", {}).get("rental_id")
		if rental_id:
			try:
				rental = Rental.objects.get(id=rental_id)
				rental.is_paid = True
				rental.status = "OWNED"
				rental.save()
			except Rental.DoesNotExist:
				pass
	elif event["type"] == "payment_intent.payment_failed":
		intent = event["data"]["object"]
		rental_id = intent.get("metadata", {}).get("rental_id")
		if rental_id:
			try:
				rental = Rental.objects.get(id=rental_id)
				rental.is_paid = False
				rental.status = "REJECTED"
				rental.save()
			except Rental.DoesNotExist:
				pass

	return HttpResponse(status=200)


def search_suggestions(request):
	q = request.GET.get('q', '').strip()
	cars = Car.objects.none()
	if q:
		cars = Car.objects.filter(
			Q(brand__icontains=q) | Q(model__icontains=q)
		)[:5]

	data = [{
		"id": c.id,
		"title": str(c),
		"image": c.image.url,
		"price": c.price_per_day,
		"url": c.get_absolute_url(),
	} for c in cars]

	return JsonResponse(data, safe=False)


@login_required
def notifications_dropdown(request):
	qs = request.user.notifications.all()[:6]
	return render(request, "partials/notifications_dropdown.html",
	              {"notifications": qs})


@login_required
def notifications_page(request):
	qs = request.user.notifications.all()
	qs.filter(is_read=False).update(is_read=True)
	return render(request, "notifications.html",
	              {"notifications": qs})


def dealer_detail(request, pk):
	dealer = get_object_or_404(Dealer, pk=pk)
	return render(request, "dealers_detail.html", {"dealer": dealer})


@login_required
@superuser_required
def admin_dealers_list(request):
	"""
	Список всех дилеров для суперпользователя.
	"""
	dealers = Dealer.objects.order_by('name')
	return render(request, 'admin/dealers_list.html', {
		'dealers': dealers,
	})


@login_required
@superuser_required
def admin_dealer_detail(request, pk):
	"""
	Детальная страница дилера:
	- основная информация,
	- список машин,
	- список аренд,
	- суммарный доход.
	"""
	dealer = get_object_or_404(Dealer, pk=pk)
	cars = Car.objects.filter(dealer=dealer)
	rentals = Rental.objects.filter(car__dealer=dealer).select_related('car', 'user')
	total_revenue = rentals.filter(status='COMPLETED') \
		                .aggregate(sum=Sum('total_price'))['sum'] or 0

	return render(request, 'admin/dealer_detail.html', {
		'dealer': dealer,
		'cars': cars,
		'rentals': rentals,
		'total_revenue': total_revenue,
	})


def car_reviews_view(request, car_id):
	"""
	Отображает страницу со всеми отзывами для конкретного автомобиля.
	"""
	car = get_object_or_404(Car, id=car_id)
	reviews = CarReview.objects.filter(car=car).order_by('-created_at')

	context = {
		'car': car,
		'reviews': reviews
	}
	return render(request, 'car_reviews.html', context)
